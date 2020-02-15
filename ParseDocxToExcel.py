#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import os
import random
import string

import xlrd
import xlsxwriter
from docx import Document


#读取指定docx文档中的所有表格
#输入参数：docxFilename - 输入的docx文档名称
#输入参数：specText - docx文档中指定的目录名称
#输出参数：xtDataSheet - xlsx文档中的Sheet名称
def readSpecTable(docxFilename, specText, xtDataSheet):

    #定义每个表格所属的业务域
    leibie = ''
    # 是否找到的第一个表格
    isFirstTable = True
    #计数器
    rowCounter = 0
    # 打开word文档
    document = Document(docxFilename)
    # 获取打开的word文档中所有的段落
    paragraphs = document.paragraphs
    # 获取打开的word文档中所有的表格
    allTables = document.tables
    print('    >>> 发现表格 ' + str(len(allTables)) + ' 个！')
    # 以防出现乱码，转为UTF-8类型
    specText = specText.encode('utf-8').decode('utf-8')
    # 遍历所有的段落
    for aPara in paragraphs:
        tableNumber = 1
        #类型为“Heading 2”为业务域，即标题二
        if aPara.style.name == 'Heading 2':
            leibie = aPara.text
        if aPara.text == specText:
            ele = aPara._p.getnext()
            while (ele.tag != '' and ele.tag[-3:] != 'tbl'):
                ele = ele.getnext()
            if ele.tag != '':
                for aTable in allTables:
                    print('    >>> 正在处理第' + str(tableNumber) + '个表格，请稍后......')
                    tableNumber = tableNumber + 1
                    print('        >>> 一共 ' + str(len(aTable.rows)) + ' 行，' +
                          str(len(aTable.columns)) + ' 列！')
                    if aTable._tbl == ele:
                        # 行
                        for i in range(len(aTable.rows)):
                            if i != 1:
                                # 列
                                if i != 0:
                                    if isFirstTable:
                                        xtDataSheet.write(
                                            rowCounter + i, 0, leibie)
                                    else:
                                        xtDataSheet.write(
                                            rowCounter + i - 1, 0, leibie)
                                for j in range(len(aTable.columns)):
                                    print("            >>>正在写入 [" + str(i) +
                                          "," + str(j) + "] - " +
                                          aTable.cell(i, j).text)
                                    if isFirstTable:
                                        xtDataSheet.write(
                                            rowCounter + i,
                                            j + 1,
                                            # 第i行第j列写入字符串
                                            aTable.cell(i, j).text)
                                    else:
                                        if i != 0:
                                            xtDataSheet.write(
                                                rowCounter + i - 1, j + 1,
                                                aTable.cell(
                                                    i, j).text)  # 第i行第j列写入字符串
                        # 计数器
                        rowCounter = len(aTable.rows)
                        isFirstTable = False
    return xtDataSheet


#读取指定目录下所有的docx文档
#输入参数： 给的包含docx文档的目录
#返回： 包含docx文件名的列表
def readFiles(dirName):
    #定义输出符合条件的文件名称
    outFiles = []
    # 读取指定目录下的所有文件

    dirs = os.listdir(dirName)
    # 循环读取路径下的文件并筛选输出
    for dir in dirs:
        # 筛选docx文件，忽略其他类型的文件
        if os.path.splitext(dir)[1] == ".docx":
            outFiles.append(dir)
    return outFiles


#合并xlsx
def mergeXls(inXls, outXls):
    isFirstSheet = True
    # 读取数据
    xtData = []
    wb = xlrd.open_workbook(inXls)

    for sheet in wb.sheets():
        for rownum in range(sheet.nrows):
            #处理标题行
            if isFirstSheet:
                xtData.append(sheet.row_values(rownum))
                isFirstSheet = False
            else:
                if rownum != 0:
                    xtData.append(sheet.row_values(rownum))
    # 写入数据
    workbook = xlsxwriter.Workbook(outXls)
    worksheet = workbook.add_worksheet('界面要素')
    font = workbook.add_format({"font_size": 10})
    for i in range(len(xtData)):
        for j in range(len(xtData[i])):
            worksheet.write(i, j, xtData[i][j], font)
    # 关闭文件流
    workbook.close()

#根据数据字典，
def finalXlsx(inXls1, inXls2, outXls):
    #存放相同的数据
    xtData = []
    #存放不同的数据
    btData = []
    wb1 = xlrd.open_workbook(inXls1)  #打开文件
    sheet1 = wb1.sheet_by_index(0)  #通过xuhao获取表格
    wb2 = xlrd.open_workbook(inXls2)  #打开文件
    sheet2 = wb2.sheet_by_name('数据项清单')  #通过名字获取表格
    xtData.append(sheet1.row_values(0) + sheet2.row_values(1))
    for rownum1 in range(sheet1.nrows):
        isSearch = False
        for rownum2 in range(sheet2.nrows):
            if sheet1.cell(rownum1, 5).value == sheet2.cell(rownum2, 0).value:
                print('    >>> 找到编号为[' + sheet2.cell(rownum2, 0).value +
                      ']的记录，正在合并中......')
                xtData.append(
                    sheet1.row_values(rownum1) + sheet2.row_values(rownum2))
                isSearch = True
        if not isSearch:
            print('    >>> 发现数据字典中没有的项，正在整理到新的Sheet')
            strValue = sheet1.row_values(rownum1)
            #忽略空行
            print('    >>> ' + str(strValue) + ' - ' + str(len(str(strValue))))
            if len(str(strValue)) != 44:
                btData.append(strValue)

    # 写入和数据字段匹配的数据
    workbook = xlsxwriter.Workbook(outXls)
    worksheet = workbook.add_worksheet('数据字典-ID')
    font = workbook.add_format({"font_size": 10})
    for i in range(len(xtData)):
        for j in range(len(xtData[i])):
            worksheet.write(i, j, xtData[i][j], font)
    #写入和数据字典不匹配的文件
    worksheet = workbook.add_worksheet('差异数据')
    font = workbook.add_format({"font_size": 9})
    #TODO: 需要处理Excel中的空行
    for i in range(len(btData)):
        for j in range(len(btData[i])):
            worksheet.write(i, j, btData[i][j], font)
    # 关闭文件流
    workbook.close()


# 主函数
if __name__ == '__main__':
    #存放待解析的Word文档目录名
    rootDir = 'D:\\workspace\\PythonTester\\'
    inputDirName = rootDir + "inputfiles\\"
    xlsxDuoName = rootDir + 'outfiles\\界面要素（多Sheet）.xlsx'
    xlsxDanName = rootDir + 'outfiles\\界面要素（单Sheet）.xlsx'
    xlsxDataDict = rootDir + 'inputfiles\\数据字典.xlsx'
    xlsxFinal = rootDir + 'outfiles\\界面要素（ID）.xlsx'
    # 打开excel，如果有此文件，则覆盖
    writebook = xlsxwriter.Workbook(xlsxDuoName)
    #生成一个新的Sheet，名称为“要素清单”
    docsFilenames = readFiles(inputDirName)
    for docsFilename in docsFilenames:
        print("正在处理文件[" + docsFilename + "], 请稍后......")
        #----2019.11.27增加内容------
        #如果文件名，即docsFilename大于31个字符，则截取
        sheetName = docsFilename
        if len(docsFilename) > 30:
            #生成随机字符串，用于新的SheetName
            ranStr = ''.join(
                random.sample(string.ascii_letters + string.digits, 8))
            sheetName = docsFilename[0:10] + ranStr
        xtDataSheet = writebook.add_worksheet(sheetName)
        #-----------------------
        allDocxFilename = inputDirName + docsFilename
        readSpecTable(allDocxFilename, '界面要素', xtDataSheet)
        print('处理完毕！！！')
    # 设置打开的Docx文档
    #docxName = 'D:\\Python\\客户信息管理应用审批客户信息修改用例说明书1.docx'
    #设置要写入的Excel文档，注意扩展名必须为xls
    print('正在写入' + xlsxDuoName + ",请稍后......")
    writebook.close()
    print('写入完毕！！！')
    print('正在合并' + xlsxDuoName + '为' + xlsxDanName + ', 请稍后......')
    mergeXls(xlsxDuoName, xlsxDanName)
    print('合并完成！！！')
    print('正在根据数据字典文件' + xlsxDataDict + '生成最终的' + xlsxFinal + '文件, 请稍后......')
    finalXlsx(xlsxDanName, xlsxDataDict, xlsxFinal)
    print('处理完成，谢谢！！！')
